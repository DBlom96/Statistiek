from docx import Document
from datetime import datetime
import comtypes.client

# Load template
doc = Document("FMW_titelblad.docx")

vaknaam = input("Vaknaam:")
vakcode = input("Vakcode:")
datum = input("Datum:")
examinator = input("Examinator:")
aantal_opgaven = input("Aantal opgaven:")
aantal_paginas = input("Aantal pagina's:")
tentamentijd = input("Tijdsduur:")

doc_path = r"C:\Users\Admin\Documents\GitHub\Statistiek\{FMW_titelblad_" + datum + ".docx"
pdf_path = r"C:\Users\Admin\Documents\GitHub\Statistiek\FMW_titelblad_" + datum + ".pdf"


# Replace placeholders
replacements = {
    "{{vaknaam}}": vaknaam,
    "{{datum}}": datum,
    "{{examinator_naam}}": examinator,
    "{{num_opgaves}}": aantal_opgaven,
    "{{vakcode}}": vakcode,
    "{{tentamentijd}}": tentamentijd,
    "{{num_paginas}}": aantal_paginas
}

def replace_text(paragraphs):
    for para in paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

replace_text(doc.paragraphs)

# Replace text inside table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_text(cell.paragraphs)

# Save the filled document
doc.save(doc_path)

def convert_docx_to_pdf(input_path, output_path):
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=17)  # 17 = wdFormatPDF
    doc.Close()
    word.Quit()

convert_docx_to_pdf(doc_path, pdf_path)
