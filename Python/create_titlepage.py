from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import comtypes.client
import locale

# Load template
doc = Document(fr"C:\Users\Admin\Documents\GitHub\Statistiek\Python\Tentamenvoorblad.docx")
locale.setlocale(locale.LC_TIME, "nl_NL.UTF-8")

vaknaam = "Statistiek deel 2 (tweede kans)"
vakcode = "STA#2"
datum = "20251113"
tentamentijd = "9:00-12:00"
examinator = "Dr. ir. D.A.M.P. Blom"
peer_review = "Dr. ir. B. Westerweel"
aantal_opgaven = "4"
aantal_paginas = "4"

formatted_date = datetime.strptime(datum, "%Y%m%d").strftime("%d %B %Y")
if formatted_date[0] == "0":
    formatted_date = formatted_date[1:]
print(formatted_date)

doc_path = fr"C:\Users\Admin\Documents\GitHub\Statistiek\Python\FMW_titelblad_{datum}.docx"
pdf_path = fr"C:\Users\Admin\Documents\GitHub\Statistiek\Python\FMW_titelblad_{datum}.pdf"

# Replace placeholders
replacements = {
    "{{vaknaam}}": vaknaam,
    "{{datum}}": str(formatted_date),
    "{{examinator_naam}}": examinator,    
    "{{num_paginas}}": aantal_paginas,
    "{{vakcode}}": vakcode,
    "{{tentamentijd}}": tentamentijd,
    "{{peer_review}}": peer_review,
    "{{num_opgaves}}": aantal_opgaven
}

def replace_text(paragraphs):
    for para in paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
                para.alignment = WD_ALIGN_VERTICAL.TOP # Align text to center

                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(14)

replace_text(doc.paragraphs)

# Replace text inside table cells and center-align
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_text(cell.paragraphs)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

# Save the filled document
doc.save(doc_path)

def convert_docx_to_pdf(input_path, output_path):
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(input_path)
    doc.SaveAs(output_path, FileFormat=17)  # 17 = wdFormatPDF
    doc.Close()
    word.Quit()

convert_docx_to_pdf(doc_path, pdf_path)
